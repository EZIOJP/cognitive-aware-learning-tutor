"""Export lecture notes (markdown) to PDF or Word with embedded images."""

from __future__ import annotations

import base64
import io
import re
from html import escape
from pathlib import Path
from typing import Literal

from backend.paths import NOTES_DIR, SNAPSHOTS_DIR
from backend.transcripts.snapshots import resolve_snapshot_path

ExportFormat = Literal["pdf", "docx"]

IMAGE_MD_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
SNAPSHOT_API_RE = re.compile(r"/api/transcripts/snapshots/([^/]+)/(\d+)\.png", re.I)

PDF_CSS = """
@page { size: A4; margin: 2cm; }
body { font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.55; color: #111; }
h1 { font-size: 22pt; margin: 0 0 12pt; border-bottom: 1px solid #ccc; padding-bottom: 6pt; }
h2 { font-size: 16pt; margin: 18pt 0 8pt; color: #0f5132; }
h3 { font-size: 13pt; margin: 14pt 0 6pt; }
p { margin: 0 0 8pt; }
ul, ol { margin: 0 0 10pt 18pt; }
li { margin-bottom: 4pt; }
blockquote { border-left: 3px solid #10b981; margin: 10pt 0; padding: 6pt 12pt; background: #f4f4f5; font-style: italic; }
pre { background: #1e1e1e; color: #e8e8e8; padding: 10pt; border-radius: 4pt; font-size: 9pt;
      white-space: pre-wrap; word-wrap: break-word; margin: 10pt 0; }
code { font-family: Consolas, monospace; font-size: 9pt; }
table { border-collapse: collapse; width: 100%; margin: 10pt 0; }
th, td { border: 1px solid #ccc; padding: 6pt 8pt; text-align: left; }
th { background: #f0f0f0; }
img { max-width: 100%; height: auto; margin: 10pt 0; display: block; }
.mermaid-label { font-size: 9pt; color: #666; margin-bottom: 4pt; }
.page-break { page-break-before: always; }
"""


def _folder_from_relative(rel: str) -> str:
    parts = rel.replace("\\", "/").split("/")
    return "/".join(parts[:-1]) if len(parts) > 1 else ""


def resolve_image_path(src: str, note_relative: str) -> Path | None:
    src = (src or "").strip()
    if not src or src.startswith("placeholder"):
        return None

    m = SNAPSHOT_API_RE.search(src)
    if m:
        try:
            path = resolve_snapshot_path(m.group(1), int(m.group(2)))
            return path if path.is_file() else None
        except (ValueError, OSError):
            return None

    notes_root = NOTES_DIR.resolve()
    folder = _folder_from_relative(note_relative)

    candidates: list[Path] = []
    if src.startswith("/"):
        candidates.append((notes_root / src.lstrip("/")).resolve())
    else:
        candidates.append((notes_root / src).resolve())
        if folder:
            candidates.append((notes_root / folder / src).resolve())

    snap_match = re.match(r"snapshots/([^/]+)/(\d+)\.png", src.replace("\\", "/"), re.I)
    if snap_match:
        try:
            candidates.append(resolve_snapshot_path(snap_match.group(1), int(snap_match.group(2))))
        except ValueError:
            pass

    for path in candidates:
        try:
            if path.is_file() and path.is_relative_to(notes_root):
                return path
        except ValueError:
            continue
        try:
            if path.is_file() and path.is_relative_to(SNAPSHOTS_DIR.resolve()):
                return path
        except ValueError:
            continue
    return None


def _image_to_data_uri(path: Path) -> str:
    ext = path.suffix.lower()
    mime = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif", ".webp": "image/webp"}.get(
        ext, "image/png"
    )
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{data}"


def _split_blocks(content: str) -> list[tuple[str, str]]:
    """Split markdown into (kind, payload) blocks."""
    blocks: list[tuple[str, str]] = []
    lines = content.replace("\r\n", "\n").split("\n")
    i = 0
    para_buf: list[str] = []

    def flush_para() -> None:
        nonlocal para_buf
        if para_buf:
            blocks.append(("paragraph", "\n".join(para_buf).strip()))
            para_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_para()
            lang = stripped[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("fence", f"{lang}\n" + "\n".join(code_lines).rstrip()))
            if i < len(lines):
                i += 1
            continue

        if stripped in ("---", "***", "___"):
            flush_para()
            blocks.append(("hr", ""))
            i += 1
            continue

        hm = re.match(r"^(#{1,6})\s+(.+)$", line)
        if hm:
            flush_para()
            blocks.append(("heading", f"{len(hm.group(1))}|{hm.group(2).strip()}"))
            i += 1
            continue

        im = IMAGE_MD_RE.match(stripped)
        if im:
            flush_para()
            blocks.append(("image", f"{im.group(1)}|{im.group(2)}"))
            i += 1
            continue

        if not stripped:
            flush_para()
            i += 1
            continue

        para_buf.append(line)
        i += 1

    flush_para()
    return blocks


def _markdown_tables_and_lists_to_html(text: str) -> str:
    """Light GFM: paragraphs, lists, blockquotes, inline code/bold."""
    import markdown as md_lib

    return md_lib.markdown(
        text,
        extensions=["tables", "nl2br", "sane_lists"],
        output_format="html5",
    )


def build_export_html(content: str, *, title: str, note_relative: str) -> str:
    parts: list[str] = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'/>",
        f"<title>{escape(title)}</title>",
        f"<style>{PDF_CSS}</style></head><body>",
        f"<h1>{escape(title)}</h1>",
    ]

    for kind, text in _split_blocks(content):
        if kind == "hr":
            parts.append("<hr/>")
        elif kind == "heading":
            level, title_text = text.split("|", 1)
            tag = f"h{min(int(level), 6)}"
            parts.append(f"<{tag}>{escape(title_text)}</{tag}>")
        elif kind == "fence":
            lang_line, code = text.split("\n", 1)
            lang = lang_line.strip()
            label = "Mermaid diagram (source)" if lang == "mermaid" else f"{lang or 'code'}"
            parts.append(f'<p class="mermaid-label">{escape(label)}</p>')
            parts.append(f"<pre><code>{escape(code)}</code></pre>")
        elif kind == "image":
            alt, src = text.split("|", 1)
            path = resolve_image_path(src, note_relative)
            if path:
                uri = _image_to_data_uri(path)
                parts.append(f'<img src="{uri}" alt="{escape(alt)}"/>')
                if alt:
                    parts.append(f'<p><em>{escape(alt)}</em></p>')
            else:
                parts.append(f'<p><em>[Image: {escape(alt or "placeholder")}]</em></p>')
        elif kind == "paragraph":
            parts.append(_markdown_tables_and_lists_to_html(text))

    parts.append("</body></html>")
    return "".join(parts)


def export_note_pdf(content: str, *, title: str, note_relative: str) -> bytes:
    try:
        from xhtml2pdf import pisa
    except ImportError as exc:
        raise RuntimeError("PDF export requires xhtml2pdf: pip install xhtml2pdf markdown python-docx") from exc

    html = build_export_html(content, title=title, note_relative=note_relative)
    buf = io.BytesIO()
    status = pisa.CreatePDF(html, dest=buf, encoding="utf-8")
    if status.err:
        raise RuntimeError("PDF generation failed")
    return buf.getvalue()


def export_note_docx(content: str, *, title: str, note_relative: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("Word export requires python-docx: pip install python-docx") from exc

    doc = Document()
    doc.add_heading(title, level=0)

    for kind, text in _split_blocks(content):
        if kind == "hr":
            doc.add_paragraph("—" * 40)
        elif kind == "heading":
            level, title_text = text.split("|", 1)
            doc.add_heading(title_text, level=min(int(level), 4))
        elif kind == "fence":
            lang_line, code = text.split("\n", 1)
            lang = lang_line.strip()
            label = "Mermaid diagram" if lang == "mermaid" else (lang or "Code")
            p = doc.add_paragraph()
            run = p.add_run(f"{label}:")
            run.bold = True
            code_p = doc.add_paragraph(code)
            for run in code_p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        elif kind == "image":
            alt, src = text.split("|", 1)
            path = resolve_image_path(src, note_relative)
            if path:
                try:
                    doc.add_picture(str(path), width=Inches(5.5))
                except Exception:
                    doc.add_paragraph(f"[Could not embed image: {alt or path.name}]")
                if alt:
                    cap = doc.add_paragraph(alt)
                    cap.runs[0].italic = True
            else:
                doc.add_paragraph(f"[Image placeholder: {alt or 'no file'}]")
        elif kind == "paragraph":
            plain_lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
            for line in plain_lines:
                if line.startswith("> "):
                    doc.add_paragraph(line[2:], style="Intense Quote")
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif re.match(r"^\d+\.\s", line):
                    doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
                else:
                    doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_note(
    content: str,
    *,
    title: str,
    note_relative: str,
    fmt: ExportFormat,
) -> tuple[bytes, str, str]:
    """Return (bytes, media_type, filename)."""
    safe = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")[:60] or "lecture_notes"
    if fmt == "pdf":
        data = export_note_pdf(content, title=title, note_relative=note_relative)
        return data, "application/pdf", f"{safe}.pdf"
    data = export_note_docx(content, title=title, note_relative=note_relative)
    return (
        data,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{safe}.docx",
    )


def combine_folder_notes(paths: list[tuple[str, str, str]]) -> tuple[str, str]:
    """Merge (relative_path, title, content) list into one document."""
    if not paths:
        raise ValueError("No notes to export.")
    parts: list[str] = []
    folder_title = _folder_from_relative(paths[0][0]) or "Study Library"
    for rel, title, body in paths:
        parts.append(f"\n\n---\n\n# {title}\n\n{body.strip()}\n")
    combined = "\n".join(parts).strip()
    return combined, folder_title.replace("/", " — ")
